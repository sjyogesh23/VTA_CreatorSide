import streamlit as st
from docx import Document
import tempfile
import os

@st.cache_data
def generate_title(summary):
    response = st.session_state.co.generate(
        model='command-xlarge-nightly',
        prompt=f"Generate a title for the following summary without 'Here is a potential title:':\n\n{summary}.",
        max_tokens=20,
        temperature=0.3,
        k=0,
        stop_sequences=["--"]
    )
    return response.generations[0].text.strip()

@st.cache_data
def generate_notes(summary):
    response = st.session_state.co.generate(
        model='command-light-nightly',
        prompt=f"Generate educational notes with bullet points from the following summary:\n\n{summary}",
        temperature=0.75,
        k=0,
        stop_sequences=["--"]
    )
    return response.generations[0].text.strip()

@st.cache_data
def save_notes_to_word(notes, title):
    doc = Document()
    doc.add_heading(title, 0)
    for line in notes.split('\n'):
        doc.add_paragraph(line)
    doc_path = os.path.join(tempfile.gettempdir(), 'generated_notes.docx')
    doc.save(doc_path)
    return doc_path


def resources():
    with st.spinner('Generating title...'):
        st.session_state.title = generate_title(st.session_state.summary)
    title = st.text_input("Title:", st.session_state.title)
    st.session_state.title = title
            
    with st.spinner('Generating notes...'):
            st.session_state.notes = generate_notes(st.session_state.summary)
    notes = st.text_area("Notes:", st.session_state.notes, height=300)
    st.session_state.notes = notes

    uploaded_files = st.file_uploader("Upload additional files (optional)", accept_multiple_files=True)
    if uploaded_files:
        st.session_state.uploaded_files = uploaded_files
            
    if st.button("Submit Resources"):
        st.session_state.notes_doc_path = save_notes_to_word(st.session_state.notes, st.session_state.title)
        st.session_state.step = 4